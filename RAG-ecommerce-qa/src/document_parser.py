"""
文档解析模块 - 支持 PDF、Word、TXT 文件
提供中英文文本处理
"""

import os
import re
from typing import List, Dict, Any
from pathlib import Path

import PyPDF2
from docx import Document
import chardet


class DocumentParser:
    """支持多格式的文档解析器"""

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        """
        初始化文档解析器

        Args:
            chunk_size: 每块最大字符数
            chunk_overlap: 块间重叠字符数
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        解析文档并返回结构化内容

        Args:
            file_path: 文档文件路径

        Returns:
            包含解析内容和元数据的字典
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}")

        # 根据文件类型提取文本
        if ext == '.pdf':
            text = self._parse_pdf(file_path)
        elif ext in ('.docx', '.doc'):
            text = self._parse_docx(file_path)
        elif ext == '.txt':
            text = self._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        # 清洗和规范化文本
        text = self._clean_text(text)

        # 分块
        chunks = self._split_text(text)

        return {
            'file_name': path.name,
            'file_type': ext,
            'file_size': path.stat().st_size,
            'total_chars': len(text),
            'chunks': chunks,
            'chunk_count': len(chunks)
        }

    def _parse_pdf(self, file_path: str) -> str:
        """从 PDF 文件提取文本"""
        text_parts = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return '\n'.join(text_parts)

    def _parse_docx(self, file_path: str) -> str:
        """从 Word 文档提取文本（段落 + 表格）"""
        doc = Document(file_path)
        text_parts = []

        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # 提取表格文本（每行用换行连接，列之间用分隔符）
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return '\n'.join(text_parts)

    def _parse_txt(self, file_path: str) -> str:
        """从 TXT 文件提取文本（自动检测编码）"""
        # 读取原始字节用于编码检测
        with open(file_path, 'rb') as f:
            raw_data = f.read()

        # 检测编码
        result = chardet.detect(raw_data)
        encoding = result.get('encoding', 'utf-8') or 'utf-8'

        # 尝试检测到的编码，失败则尝试常见编码
        encodings_to_try = [encoding, 'utf-8', 'gbk', 'gb2312', 'latin-1']
        for enc in encodings_to_try:
            try:
                return raw_data.decode(enc)
            except (UnicodeDecodeError, LookupError):
                continue

        # 最后手段：忽略错误解码
        return raw_data.decode('utf-8', errors='ignore')

    def _clean_text(self, text: str) -> str:
        """清洗和规范化文本内容"""
        # 统一换行符（Windows \r\n → Unix \n）
        text = text.replace('\r\n', '\n')
        # 去除多余空白（保留段落结构）
        text = re.sub(r'[ \t]+', ' ', text)
        # 规范化换行
        text = re.sub(r'\n{3,}', '\n\n', text)
        # 去除可能干扰处理的特殊字符
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
        return text.strip()

    def _split_text(self, text: str) -> List[Dict[str, Any]]:
        """
        将文本分割为重叠的块

        Returns:
            块字典列表（包含文本和元数据）
        """
        if not text:
            return []

        # 检测是否为对话数据集格式（每行以 "1\t" 或 "1 " 开头）
        lines = text.split('\n')
        dialogue_lines = [l.strip() for l in lines if l.strip().startswith('1\t') or l.strip().startswith('1 ')]

        # 如果对话行数 >= 3，按对话行分块（每条对话一个块）
        if len(dialogue_lines) >= 3:
            return self._split_dialogues(text)

        chunks = []
        # 优先按段落分块
        paragraphs = text.split('\n\n')

        current_chunk = ""
        chunk_index = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # 如果添加当前段落超过块大小，保存当前块
            if current_chunk and len(current_chunk) + len(para) + 2 > self.chunk_size:
                chunks.append({
                    'text': current_chunk,
                    'index': chunk_index,
                    'char_count': len(current_chunk)
                })
                chunk_index += 1
                # 保留上一块末尾的重叠部分
                if self.chunk_overlap > 0 and len(current_chunk) > self.chunk_overlap:
                    current_chunk = current_chunk[-self.chunk_overlap:]
                else:
                    current_chunk = ""

            # 处理超大段落
            if len(para) > self.chunk_size:
                # 按句子分割大段落
                sentences = re.split(r'([。！？.!?])', para)
                for i in range(0, len(sentences), 2):
                    sentence = sentences[i]
                    if i + 1 < len(sentences):
                        sentence += sentences[i + 1]

                    if current_chunk and len(current_chunk) + len(sentence) > self.chunk_size:
                        chunks.append({
                            'text': current_chunk,
                            'index': chunk_index,
                            'char_count': len(current_chunk)
                        })
                        chunk_index += 1
                        if self.chunk_overlap > 0 and len(current_chunk) > self.chunk_overlap:
                            current_chunk = current_chunk[-self.chunk_overlap:]
                        else:
                            current_chunk = ""

                    current_chunk += sentence
            else:
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para

        # 保存最后一块
        if current_chunk.strip():
            chunks.append({
                'text': current_chunk,
                'index': chunk_index,
                'char_count': len(current_chunk)
            })

        return chunks

    def _split_dialogues(self, text: str) -> List[Dict[str, Any]]:
        """
        按对话行分块（适用于对话数据集格式）

        每条对话以 "1\t" 开头，作为独立的知识单元。
        """
        chunks = []
        chunk_index = 0
        lines = text.split('\n')

        current_chunk = ""
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # 遇到新的对话行，保存当前块，开始新块
            if stripped.startswith('1\t') or stripped.startswith('1 '):
                if current_chunk.strip():
                    chunks.append({
                        'text': current_chunk.strip(),
                        'index': chunk_index,
                        'char_count': len(current_chunk.strip())
                    })
                    chunk_index += 1
                current_chunk = stripped
            else:
                # 非对话行（标题、分隔线等），追加到当前块
                if current_chunk:
                    current_chunk += "\n" + stripped
                else:
                    current_chunk = stripped

        # 保存最后一块
        if current_chunk.strip():
            chunks.append({
                'text': current_chunk.strip(),
                'index': chunk_index,
                'char_count': len(current_chunk.strip())
            })

        return chunks
